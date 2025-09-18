import io
import logging
from typing import Dict, Any, Optional
from fastapi import UploadFile, HTTPException
import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)

class DocumentProcessingService:
    """Service for processing various document types and extracting text content"""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self._extract_pdf_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'application/msword': self._extract_doc_text,
            'text/plain': self._extract_txt_text,
            'text/csv': self._extract_txt_text,
        }
    
    async def process_document(self, file: UploadFile) -> Dict[str, Any]:
        """
        Process uploaded document and extract text content
        
        Args:
            file: Uploaded file object
            
        Returns:
            Dict containing extracted text and metadata
        """
        try:
            # Validate file
            if not file.filename:
                raise HTTPException(status_code=400, detail="No filename provided")
            
            # Read file content
            file_content = await file.read()
            
            # Detect file type
            file_type = self._detect_file_type(file_content, file.filename)
            
            if file_type not in self.supported_types:
                raise HTTPException(
                    status_code=400, 
                    detail=f"Unsupported file type: {file_type}. Supported types: PDF, DOCX, DOC, TXT, CSV"
                )
            
            # Extract text based on file type
            extractor = self.supported_types[file_type]
            text_content = await extractor(file_content)
            
            if not text_content or not text_content.strip():
                raise HTTPException(status_code=400, detail="No text content found in document")
            
            # Prepare response
            result = {
                "filename": file.filename,
                "file_type": file_type,
                "file_size": len(file_content),
                "content": text_content.strip(),
                "content_length": len(text_content.strip()),
                "status": "success"
            }
            
            logger.info(f"Successfully processed document: {file.filename} ({file_type}, {len(text_content)} chars)")
            return result
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Error processing document {file.filename}: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to process document: {str(e)}")
    
    def _detect_file_type(self, file_content: bytes, filename: str) -> str:
        """Detect file type using filename extension and basic file signatures"""
        try:
            # Get file extension
            extension = filename.lower().split('.')[-1] if '.' in filename else ''
            
            # Extension-based detection
            extension_map = {
                'pdf': 'application/pdf',
                'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                'doc': 'application/msword',
                'txt': 'text/plain',
                'csv': 'text/csv'
            }
            
            # Check if extension is supported
            if extension in extension_map:
                return extension_map[extension]
            
            # Try basic file signature detection for PDF
            if file_content.startswith(b'%PDF'):
                return 'application/pdf'
            
            # Try basic file signature detection for DOCX (ZIP-based)
            if file_content.startswith(b'PK\x03\x04') and b'word/' in file_content[:1024]:
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
            # Default to text if no specific type detected
            return 'text/plain'
            
        except Exception as e:
            logger.warning(f"File type detection failed for {filename}: {str(e)}")
            # Fall back to text/plain
            return 'text/plain'
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content += page.extract_text() + "\n"
            
            return text_content
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + " "
                    text_content += "\n"
            
            return text_content
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Failed to extract text from DOCX: {str(e)}")
    
    async def _extract_doc_text(self, file_content: bytes) -> str:
        """Extract text from DOC file (legacy Word format)"""
        # Note: python-docx doesn't support .doc files, only .docx
        # For .doc files, we would need python-docx2txt or similar
        raise HTTPException(
            status_code=400, 
            detail="Legacy .doc files are not supported. Please convert to .docx format."
        )
    
    async def _extract_txt_text(self, file_content: bytes) -> str:
        """Extract text from TXT/CSV file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            return file_content.decode('utf-8', errors='replace')
            
        except Exception as e:
            logger.error(f"Error extracting TXT text: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Failed to extract text from text file: {str(e)}")
